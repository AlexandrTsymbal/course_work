import re

from docx2pdf import convert
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION

import pandas as pd


class StudyPlan:
    __all_data = []
    basic_info = []
    subjects = []
    extra = []
    file_name = ''

    def __init__(self, file: str, sheet: str):
        self.file_name = file
        df = pd.read_excel(file, sheet_name=sheet)

        for index, row in df.iterrows():
            row_data = row.values.tolist()
            self.__all_data.append(row_data)

        self.__create_info()
        self.__add_extra(file)

    def __create_info(self):
        moduls = []
        self.basic_info = []
        create_data = False
        fill_list_moduls = False

        for data in self.__all_data:
            if (not create_data) and ('Шифр' in str(data[1])):
                create_data = True

            if create_data and ('Стандарт ВО' not in str(data[1])):
                self.basic_info.append(data[1])
                if str(data[3]) != 'nan':
                    self.basic_info.append(data[3])

            if create_data and ('Стандарт ВО' in str(data[1])):
                create_data = False
                fill_list_moduls = True

            if fill_list_moduls and str(data[0]).isdigit():
                moduls.append(data)

        moduls.pop(0)
        moduls.pop(-1)
        moduls.pop(-1)

        self.__itemizing(moduls)

    def __itemizing(self, moduls):
        for mod in moduls:
            if ('М' not in mod[1]) and ('Б' not in mod[1]):
                if int(mod[1][0]) < 4:
                    dis = {'name': mod[2], 'sems': {}, 'less': []}
                    if str(mod[3]) != 'nan':
                        dis['sems'][mod[3]] = "Экзамен"
                    if str(mod[4]) != 'nan':
                        dis['sems'][mod[4]] = "Зачет"

                    dis['less'].append(mod[14])
                    dis['less'].append(mod[15])
                    dis['less'].append(mod[16])
                    self.subjects.append(dis)

        def parse_sems(sems):
            if isinstance(sems, list):
                return [int(num) for s in sems for num in re.findall(r'\d+', str(s))]
            else:
                return [int(num) for num in re.findall(r'\d+', sems)]

    def __add_extra(self, file):
        all_data = []
        moduls = []
        df = pd.read_excel(file, sheet_name='приложение (1)')

        for index, row in df.iterrows():
            row_data = row.values.tolist()
            all_data.append(row_data)

        fl = False
        for data in all_data:
            if 'Стандарт' in str(data[1]):
                fl = True

            if fl and str(data[0]).isdigit():
                moduls.append(data)

        for mod in moduls:
            if ('М' not in mod[1]) and ('Б' not in mod[1]):
                if int(mod[1][0]) < 4:
                    dis = {'name': mod[2], 'sems': {}, 'less': []}
                    if str(mod[3]) != 'nan':
                        dis['sems'][mod[3]] = "Экзамен"
                    if str(mod[4]) != 'nan':
                        dis['sems'][mod[4]] = "Зачет"

                    dis['less'].append(mod[14])
                    dis['less'].append(mod[15])
                    dis['less'].append(mod[16])
                    self.extra.append(dis)

    def get_info(self):
        for info in self.basic_info:
            print(info + '\n')
        for dis in self.subjects:
            print(dis)
            print()
        for ex in self.extra:
            print(ex)

    def to_files(self, dir: str = '.'):

        doc = Document()
        paragraph_cent = doc.add_paragraph()
        paragraph_cent.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = paragraph_cent.add_run(f'{self.basic_info[1]}\n\n{self.basic_info[3]}\n\n{self.basic_info[5]}\n\n'
                                     f'{self.basic_info[6]}\n\n{self.basic_info[7]}\n\n{self.basic_info[8]}\n\n')
        run.bold = True
        run.font.size = Pt(24)

        doc.add_section(WD_SECTION.NEW_PAGE)
        paragraph_cent = doc.add_paragraph()
        paragraph_cent.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph_cent.add_run(f'Основные дисциплины\n')
        run.font.size = Pt(20)

        def iterr(lst):
            for i in range(1, 12):
                data = []
                subj = []
                ls = '\n'
                ind = 1

                for sb in lst:
                    for sm in sb['sems']:
                        if str(i) == sm or (len(sm) > 2 and int(sm[0]) <= i <= int(sm[2:])):

                            subj.append('\n' + str(ind) + '\n')
                            ind += 1
                            subj.append('\n' + sb['name'] + '\n')

                            if str(sb['less'][0]) != 'nan':
                                ls += 'Лекции\n'
                            if str(sb['less'][1]) != 'nan':
                                ls += 'Практики\n'
                            if str(sb['less'][2]) != 'nan':
                                ls += 'Лабораторные\n'

                            subj.append(ls)
                            subj.append('\n' + sb['sems'][sm] + '\n')

                            data.append(subj)
                            subj = []
                            ls = '\n'
                            break

                if data:
                    paragraph_cent = doc.add_paragraph()
                    paragraph_cent.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph_cent.add_run(f'{str(i)} Семестр\n')
                    run.font.size = Pt(18)

                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'
                    for j, col in enumerate(table.columns):
                        if j == 0:
                            col.width = int(914400 * 0.1)
                        elif j == 1:
                            col.width = int(914400 * 10)
                        else:
                            col.width = int(914400 * 7)

                    for cell in table.columns[0].cells:
                        cell.width = 0.5 * 1440
                    for cell in table.columns[1].cells:
                        cell.width = 1.5 * 1440
                    for cell in table.columns[2].cells:
                        cell.width = 1.5 * 1440
                    for cell in table.columns[3].cells:
                        cell.width = 1.5 * 1440

                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = '\nНомер\n'
                    hdr_cells[1].text = '\nНаименование дисциплины\n'
                    hdr_cells[2].text = '\nВиды учебной нагрузки\n'
                    hdr_cells[3].text = '\nФорма контроля\n'

                    for row_data in data:
                        row = table.add_row().cells
                        for j, text in enumerate(row_data):
                            row[j].text = str(text)

                    doc.add_section(WD_SECTION.NEW_PAGE)

        iterr(self.subjects)
        paragraph_cent = doc.add_paragraph()
        paragraph_cent.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph_cent.add_run(f'Дисциплины по выбору\n')
        run.font.size = Pt(20)
        iterr(self.extra)

        doc.save(f'{dir}/{self.file_name[:self.file_name.find('.')]}.docx')
        self.__convert_to_pdf(dir)

    def __convert_to_pdf(self, dir: str):
        convert(f'{dir}/{self.file_name[:self.file_name.find('.')]}.docx')
